#!/usr/bin/env python3
"""
Create a test DOCX file for faults/warnings testing
"""
try:
    from docx import Document
    from docx.shared import Inches
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Faults and Warnings Mapping', 0)
    
    # Add description
    doc.add_paragraph('This document contains fault and warning mappings for PLC testing.')
    
    # Create a table with fault/warning mappings
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'TAG'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Resolution'
    
    # Add sample fault entries
    fault_entries = [
        ('{[PLC]Alarm_Fault[0].0}', 'Motor Overload Fault', 'Check motor current and thermal protection'),
        ('{[PLC]Alarm_Fault[0].1}', 'Sensor Communication Error', 'Verify sensor wiring and communication'),
        ('{[PLC]Alarm_Fault[0].2}', 'Safety Interlock Fault', 'Check safety circuit and interlocks'),
        ('{[PLC]Alarm_Fault[1].0}', 'Pneumatic Pressure Low', 'Check air supply and pressure regulator'),
        ('{[PLC]Alarm_Fault[1].1}', 'Conveyor Jam Detected', 'Clear conveyor obstruction'),
    ]
    
    # Add sample warning entries
    warning_entries = [
        ('{[PLC]Alarm_Warning[0].0}', 'High Temperature Warning', 'Monitor temperature and check cooling'),
        ('{[PLC]Alarm_Warning[0].1}', 'Low Battery Warning', 'Replace battery in backup system'),
        ('{[PLC]Alarm_Warning[0].2}', 'Maintenance Due Warning', 'Schedule maintenance according to plan'),
        ('{[PLC]Alarm_Warning[1].0}', 'Network Latency Warning', 'Check network performance'),
    ]
    
    # Add all entries to table
    all_entries = fault_entries + warning_entries
    
    for tag, description, resolution in all_entries:
        row_cells = table.add_row().cells
        row_cells[0].text = tag
        row_cells[1].text = description
        row_cells[2].text = resolution
    
    # Save the document
    filename = 'test_faults.docx'
    doc.save(filename)
    print(f"✓ Created test DOCX file: {filename}")
    print(f"  - {len(fault_entries)} fault entries")
    print(f"  - {len(warning_entries)} warning entries")
    print(f"  - {len(all_entries)} total entries")
    
except ImportError:
    print("Error: python-docx not installed. Please run: pip install python-docx")
except Exception as e:
    print(f"Error creating test DOCX: {e}")